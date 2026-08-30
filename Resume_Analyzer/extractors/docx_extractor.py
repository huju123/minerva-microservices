import docx
import os
import re

def extract_text_from_docx(filename):
    if not os.path.isfile(filename):
        raise FileNotFoundError(f"Invalid path: {filename}")
    text_list = []
    text_list_tables = []
    document = docx.Document(filename)
    for paragraph in document.paragraphs:
        text = paragraph.text
        text_list.append(text)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                text_list_tables.append(text)
    text_list.extend(text_list_tables)
    result = "\n".join(text_list)
    if not result or result.isspace():
        try:
            result = extract_text_from_xml_fallback(document)     
        except:
            raise ValueError("DOCX doesn't contain extractable text")
        return result, True
        
    return result, False
                
def extract_text_from_xml_fallback(document):
    regex_pattern = r"<w:t[^>]*>(.*?)</w:t>"
    text = document.element.xml
    output = re.findall(regex_pattern, text)
    result = "\n".join(output)
    return result
    